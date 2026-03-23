"""PE portfolio batch scanner for VXIS.

Supports loading multiple targets from a CSV file and scanning them
concurrently, then aggregating results into a risk-ranked summary report.
"""

from __future__ import annotations

import asyncio
import csv
import io
import logging
from collections import Counter
from dataclasses import dataclass, field
from pathlib import Path
from typing import Callable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from vxis.config.schema import VXISConfig
from vxis.core.orchestrator import ScanOrchestrator, ScanResult

logger = logging.getLogger(__name__)

_SEVERITY_ORDER: list[str] = [
    "critical",
    "high",
    "medium",
    "low",
    "informational",
]

_SEVERITY_COLORS: dict[str, tuple[int, int, int]] = {
    "critical": (123, 44, 52),
    "high": (192, 57, 43),
    "medium": (230, 126, 34),
    "low": (46, 204, 113),
    "informational": (52, 152, 219),
}

_HEADER_BG: tuple[int, int, int] = (31, 55, 93)
_ALT_ROW_BG: tuple[int, int, int] = (242, 242, 242)

_GRADE_COLORS: dict[str, tuple[int, int, int]] = {
    "A": (46, 204, 113),
    "B": (52, 152, 219),
    "C": (230, 126, 34),
    "D": (192, 57, 43),
    "F": (123, 44, 52),
}


# ---------------------------------------------------------------------------
# Data classes
# ---------------------------------------------------------------------------


@dataclass
class BatchTarget:
    """A single scan target in a portfolio batch."""

    name: str
    domain: str
    extra_domains: list[str] = field(default_factory=list)
    notes: str = ""


@dataclass
class BatchResult:
    """Result for a single target in a portfolio batch scan."""

    target: BatchTarget
    scan_result: ScanResult | None
    error: str | None = None

    @property
    def succeeded(self) -> bool:
        """Return True when the scan completed without error."""
        return self.scan_result is not None and self.error is None


# ---------------------------------------------------------------------------
# XML cell helpers (duplicated locally to avoid import coupling)
# ---------------------------------------------------------------------------


def _set_cell_bg(cell, rgb: tuple[int, int, int]) -> None:
    hex_color = "{:02X}{:02X}{:02X}".format(*rgb)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_borders(cell) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "CCCCCC")
        tcBorders.append(border)
    tcPr.append(tcBorders)


# ---------------------------------------------------------------------------
# BatchScanner
# ---------------------------------------------------------------------------


class BatchScanner:
    """Scan multiple targets from a CSV portfolio file.

    Designed for Private Equity due-diligence workflows where dozens of
    portfolio companies must be assessed quickly.

    Args:
        config: Root VXIS configuration object shared across all scans.
    """

    def __init__(self, config: VXISConfig) -> None:
        self.config = config
        self.orchestrator = ScanOrchestrator(config)

    # ------------------------------------------------------------------
    # CSV loading
    # ------------------------------------------------------------------

    @staticmethod
    def load_targets(csv_path: Path) -> list[BatchTarget]:
        """Load scan targets from a CSV file.

        The CSV must have the following columns (header row required):
            name, domain, extra_domains, notes

        The *extra_domains* column is a semicolon-separated list of
        additional domains to associate with the target (e.g. staging
        environments or API sub-domains). It may be empty.

        Parameters
        ----------
        csv_path:
            Path to the CSV file, or a :class:`pathlib.Path` wrapping an
            ``io.StringIO`` object (used in unit tests via a file-like path).

        Returns
        -------
        list[BatchTarget]
            Ordered list of targets as read from the file.
        """
        targets: list[BatchTarget] = []

        # Support both real paths and StringIO-backed paths in tests
        if isinstance(csv_path, io.StringIO):
            reader = csv.DictReader(csv_path)
            for row in reader:
                targets.append(BatchScanner._row_to_target(row))
            return targets

        with open(csv_path, newline="", encoding="utf-8") as fh:
            reader = csv.DictReader(fh)
            for row in reader:
                targets.append(BatchScanner._row_to_target(row))

        return targets

    @staticmethod
    def _row_to_target(row: dict[str, str]) -> BatchTarget:
        """Convert a single CSV row dict to a :class:`BatchTarget`."""
        extra_raw = row.get("extra_domains", "").strip()
        extra_domains = [
            d.strip() for d in extra_raw.split(";") if d.strip()
        ]
        return BatchTarget(
            name=row.get("name", "").strip(),
            domain=row.get("domain", "").strip(),
            extra_domains=extra_domains,
            notes=row.get("notes", "").strip(),
        )

    # ------------------------------------------------------------------
    # Batch execution
    # ------------------------------------------------------------------

    async def run_batch(
        self,
        targets: list[BatchTarget],
        profile: str = "standard",
        max_concurrent: int = 3,
        on_complete: Callable[[BatchResult], None] | None = None,
    ) -> list[BatchResult]:
        """Scan all targets with a concurrency limit.

        Uses a :class:`asyncio.Semaphore` to cap the number of simultaneous
        scans. Each failed scan is captured as a :class:`BatchResult` with
        ``error`` populated rather than raising.

        Parameters
        ----------
        targets:
            Ordered list of :class:`BatchTarget` objects.
        profile:
            Scan profile name passed to :class:`ScanOrchestrator`.
        max_concurrent:
            Maximum number of concurrent scan sessions.
        on_complete:
            Optional synchronous callback invoked after each target
            completes (or fails). Receives the :class:`BatchResult`.

        Returns
        -------
        list[BatchResult]
            Results in the same order as *targets*.
        """
        semaphore = asyncio.Semaphore(max_concurrent)
        results: list[BatchResult | None] = [None] * len(targets)

        async def _scan_one(idx: int, target: BatchTarget) -> None:
            async with semaphore:
                logger.info(
                    "Batch scan [%d/%d]: starting '%s' (%s)",
                    idx + 1,
                    len(targets),
                    target.name,
                    target.domain,
                )
                try:
                    scan_result = await self.orchestrator.run_scan(
                        target=target.domain,
                        profile=profile,
                    )
                    batch_result = BatchResult(
                        target=target,
                        scan_result=scan_result,
                    )
                    logger.info(
                        "Batch scan [%d/%d]: completed '%s' — %d finding(s)",
                        idx + 1,
                        len(targets),
                        target.name,
                        len(scan_result.findings),
                    )
                except Exception as exc:  # noqa: BLE001
                    logger.warning(
                        "Batch scan [%d/%d]: failed '%s' — %s",
                        idx + 1,
                        len(targets),
                        target.name,
                        exc,
                    )
                    batch_result = BatchResult(
                        target=target,
                        scan_result=None,
                        error=str(exc),
                    )

                results[idx] = batch_result

                if on_complete is not None:
                    try:
                        on_complete(batch_result)
                    except Exception:  # noqa: BLE001
                        logger.warning(
                            "on_complete callback raised for '%s'", target.name
                        )

        # return_exceptions=True ensures that an unexpected exception escaping
        # _scan_one (e.g. from the semaphore or logging) never cancels sibling
        # tasks — each target is isolated.
        await asyncio.gather(
            *[_scan_one(i, t) for i, t in enumerate(targets)],
            return_exceptions=True,
        )

        # All slots must be populated — cast away None (guaranteed by gather)
        return [r for r in results if r is not None]  # type: ignore[misc]

    # ------------------------------------------------------------------
    # Summary report
    # ------------------------------------------------------------------

    def generate_summary_report(
        self,
        results: list[BatchResult],
        output_path: Path,
    ) -> Path:
        """Generate a PE portfolio summary DOCX report.

        The report contains:
        - Risk ranking table (A–F grade per target)
        - Portfolio-wide finding totals
        - Top-5 most common finding types
        - Recommendations prioritized by frequency

        Parameters
        ----------
        results:
            List of :class:`BatchResult` objects from :meth:`run_batch`.
        output_path:
            Destination path for the generated .docx file.

        Returns
        -------
        Path
            Resolved path of the written file.
        """
        output_path = Path(output_path).resolve()
        output_path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()
        section = doc.sections[0]
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)

        self._add_summary_cover(doc, results)
        self._add_risk_ranking_table(doc, results)
        self._add_portfolio_totals(doc, results)
        self._add_top_findings(doc, results)
        self._add_priority_recommendations(doc, results)

        doc.save(str(output_path))
        return output_path

    # ------------------------------------------------------------------
    # Internal summary section builders
    # ------------------------------------------------------------------

    def _add_summary_cover(self, doc: Document, results: list[BatchResult]) -> None:
        for _ in range(4):
            doc.add_paragraph()

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("PE Portfolio Security Assessment")
        run.font.name = "Calibri"
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = RGBColor(31, 55, 93)

        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run("Risk Ranking Summary Report")
        run2.font.size = Pt(14)
        run2.font.color.rgb = RGBColor(80, 80, 80)

        doc.add_paragraph()
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run3 = p3.add_run(
            f"{sum(1 for r in results if r.succeeded)} of {len(results)} target(s) scanned successfully"
        )
        run3.font.size = Pt(11)
        run3.font.color.rgb = RGBColor(100, 100, 100)

        # Page break after cover
        para = doc.add_paragraph()
        para.add_run().add_break(
            __import__("docx.enum.text", fromlist=["WD_BREAK"]).WD_BREAK.PAGE
        )

    def _add_risk_ranking_table(
        self, doc: Document, results: list[BatchResult]
    ) -> None:
        """Render the main risk ranking table."""
        h = doc.add_paragraph()
        run = h.add_run("1. Portfolio Risk Ranking")
        run.font.name = "Calibri"
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(31, 55, 93)
        doc.add_paragraph()

        cols = ["Company", "Domain", "Grade", "Critical", "High", "Medium", "Low", "Total", "Notes"]
        tbl = doc.add_table(rows=1, cols=len(cols))
        tbl.style = "Table Grid"

        # Header row
        for idx, col_name in enumerate(cols):
            cell = tbl.rows[0].cells[idx]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(col_name)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(9)
            _set_cell_bg(cell, _HEADER_BG)
            _set_cell_borders(cell)

        # Sort results: failed scans last, then by grade (F first for ranking)
        grade_order = {"F": 0, "D": 1, "C": 2, "B": 3, "A": 4, "?": 5}
        sorted_results = sorted(
            results,
            key=lambda r: grade_order.get(
                self.calculate_risk_grade(r.scan_result) if r.succeeded else "?",
                5,
            ),
        )

        for row_idx, result in enumerate(sorted_results):
            row = tbl.add_row()
            cells = row.cells

            target = result.target
            if result.succeeded and result.scan_result:
                counts = result.scan_result.severity_counts
                grade = self.calculate_risk_grade(result.scan_result)
                critical = counts.get("critical", 0)
                high = counts.get("high", 0)
                medium = counts.get("medium", 0)
                low = counts.get("low", 0)
                total = len(result.scan_result.findings)
                error_text = ""
            else:
                grade = "?"
                critical = high = medium = low = total = 0
                error_text = result.error or "scan failed"

            values = [
                target.name,
                target.domain,
                grade,
                str(critical),
                str(high),
                str(medium),
                str(low),
                str(total),
                error_text or target.notes,
            ]

            for col_idx, value in enumerate(values):
                cells[col_idx].paragraphs[0].clear()
                if col_idx == 2:  # Grade column
                    cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = cells[col_idx].paragraphs[0].add_run(grade)
                    rgb = _GRADE_COLORS.get(grade, (150, 150, 150))
                    run.font.color.rgb = RGBColor(*rgb)
                    run.font.bold = True
                    run.font.size = Pt(9)
                else:
                    cells[col_idx].paragraphs[0].add_run(value).font.size = Pt(9)

                if row_idx % 2 == 0:
                    _set_cell_bg(cells[col_idx], _ALT_ROW_BG)
                _set_cell_borders(cells[col_idx])

        doc.add_paragraph()

    def _add_portfolio_totals(
        self, doc: Document, results: list[BatchResult]
    ) -> None:
        """Render portfolio-wide finding totals."""
        h = doc.add_paragraph()
        run = h.add_run("2. Portfolio-Wide Totals")
        run.font.name = "Calibri"
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(31, 55, 93)

        total_counts: dict[str, int] = {s: 0 for s in _SEVERITY_ORDER}
        successful = 0
        for result in results:
            if result.succeeded and result.scan_result:
                successful += 1
                for sev, count in result.scan_result.severity_counts.items():
                    if sev in total_counts:
                        total_counts[sev] += count

        grand_total = sum(total_counts.values())

        p = doc.add_paragraph()
        p.add_run(
            f"Across {successful} successfully scanned target(s), a total of "
            f"{grand_total} finding(s) were identified."
        )

        tbl = doc.add_table(rows=1, cols=len(_SEVERITY_ORDER) + 1)
        tbl.style = "Table Grid"

        # Header
        hdr = tbl.rows[0].cells
        hdr[0].paragraphs[0].add_run("Severity").font.bold = True
        _set_cell_bg(hdr[0], _HEADER_BG)
        hdr[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        _set_cell_borders(hdr[0])
        for idx, sev in enumerate(_SEVERITY_ORDER):
            cell = hdr[idx + 1]
            run = cell.paragraphs[0].add_run(sev.capitalize())
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            _set_cell_bg(cell, _HEADER_BG)
            _set_cell_borders(cell)

        # Count row
        count_row = tbl.add_row()
        count_cells = count_row.cells
        count_cells[0].paragraphs[0].add_run("Total").font.bold = True
        _set_cell_bg(count_cells[0], _ALT_ROW_BG)
        _set_cell_borders(count_cells[0])
        for idx, sev in enumerate(_SEVERITY_ORDER):
            cell = count_cells[idx + 1]
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cell.paragraphs[0].add_run(str(total_counts[sev]))
            rgb = _SEVERITY_COLORS.get(sev, (0, 0, 0))
            run.font.color.rgb = RGBColor(*rgb)
            run.font.bold = total_counts[sev] > 0
            _set_cell_bg(cell, _ALT_ROW_BG)
            _set_cell_borders(cell)

        doc.add_paragraph()

    def _add_top_findings(self, doc: Document, results: list[BatchResult]) -> None:
        """Render top-5 most common finding types across the portfolio."""
        h = doc.add_paragraph()
        run = h.add_run("3. Top 5 Most Common Findings")
        run.font.name = "Calibri"
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(31, 55, 93)

        type_counter: Counter[str] = Counter()
        for result in results:
            if result.succeeded and result.scan_result:
                for finding in result.scan_result.findings:
                    type_counter[finding.finding_type] += 1

        top5 = type_counter.most_common(5)

        if not top5:
            doc.add_paragraph("No findings recorded across the portfolio.")
            return

        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Table Grid"

        # Header
        for idx, label in enumerate(["Rank", "Finding Type", "Occurrences"]):
            cell = tbl.rows[0].cells[idx]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(label)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(9)
            _set_cell_bg(cell, _HEADER_BG)
            _set_cell_borders(cell)

        for rank, (finding_type, count) in enumerate(top5, start=1):
            row = tbl.add_row()
            cells = row.cells
            cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cells[0].paragraphs[0].add_run(str(rank)).font.size = Pt(9)
            cells[1].paragraphs[0].add_run(finding_type).font.size = Pt(9)
            cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cells[2].paragraphs[0].add_run(str(count)).font.size = Pt(9)
            if rank % 2 == 0:
                for cell in cells:
                    _set_cell_bg(cell, _ALT_ROW_BG)
            for cell in cells:
                _set_cell_borders(cell)

        doc.add_paragraph()

    def _add_priority_recommendations(
        self, doc: Document, results: list[BatchResult]
    ) -> None:
        """Render recommendations prioritized by frequency."""
        h = doc.add_paragraph()
        run = h.add_run("4. Priority Recommendations")
        run.font.name = "Calibri"
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(31, 55, 93)

        doc.add_paragraph(
            "The following recommendations address the most frequently observed "
            "findings across the portfolio and should be prioritized accordingly."
        )

        # Collect unique remediations, weighted by how many targets have them
        remediation_counter: Counter[str] = Counter()
        for result in results:
            if result.succeeded and result.scan_result:
                seen_in_scan: set[str] = set()
                for finding in result.scan_result.findings:
                    if finding.remediation and finding.remediation not in seen_in_scan:
                        remediation_counter[finding.remediation] += 1
                        seen_in_scan.add(finding.remediation)

        top_recs = remediation_counter.most_common(10)

        if not top_recs:
            doc.add_paragraph("No specific remediation guidance available.")
            return

        for idx, (remediation, count) in enumerate(top_recs, start=1):
            p = doc.add_paragraph(style="List Number")
            run = p.add_run(remediation)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(31, 55, 93)
            # Append frequency note
            freq_run = p.add_run(f"  (affects {count} target(s))")
            freq_run.font.size = Pt(9)
            freq_run.font.color.rgb = RGBColor(120, 120, 120)
            freq_run.font.italic = True

    # ------------------------------------------------------------------
    # Grade calculation
    # ------------------------------------------------------------------

    @staticmethod
    def calculate_risk_grade(scan_result: ScanResult) -> str:
        """Derive an A–F risk grade from a scan's critical/high finding counts.

        Grading scale:
        - A: 0 critical, 0 high
        - B: 0 critical, 1–3 high
        - C: 0 critical, 4+ high  OR  1 critical
        - D: 2–3 critical
        - F: 4+ critical

        Parameters
        ----------
        scan_result:
            Completed :class:`ScanResult` instance.

        Returns
        -------
        str
            Single letter grade: 'A', 'B', 'C', 'D', or 'F'.
        """
        counts = scan_result.severity_counts
        critical = counts.get("critical", 0)
        high = counts.get("high", 0)

        if critical >= 4:
            return "F"
        if critical in (2, 3):
            return "D"
        if critical == 1 or high >= 4:
            return "C"
        if high in range(1, 4):  # 1, 2, 3
            return "B"
        # 0 critical, 0 high
        return "A"
