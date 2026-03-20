"""DOCX report generator for VXIS security assessments.

Produces NCC Group-style Word documents with cover page, executive summary,
findings table, per-finding detail sections, and methodology appendix.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from vxis.models.finding import Finding, Severity
from vxis.report.generator import ReportData

# ---------------------------------------------------------------------------
# Severity colour mapping (RGB tuples)
# ---------------------------------------------------------------------------

_SEVERITY_COLORS: dict[str, tuple[int, int, int]] = {
    "critical": (123, 44, 52),
    "high": (192, 57, 43),
    "medium": (230, 126, 34),
    "low": (46, 204, 113),
    "informational": (52, 152, 219),
}

_SEVERITY_ORDER: list[str] = [
    "critical",
    "high",
    "medium",
    "low",
    "informational",
]

# Header row fill colour (dark navy)
_HEADER_BG: tuple[int, int, int] = (31, 55, 93)
# Light grey for alternating rows
_ALT_ROW_BG: tuple[int, int, int] = (242, 242, 242)


# ---------------------------------------------------------------------------
# Low-level XML helpers
# ---------------------------------------------------------------------------


def _set_cell_bg(cell: Any, rgb: tuple[int, int, int]) -> None:
    """Set background shading colour on a table cell using direct XML manipulation."""
    hex_color = "{:02X}{:02X}{:02X}".format(*rgb)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_borders(cell: Any) -> None:
    """Apply thin border to all sides of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "CCCCCC")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _add_page_break(doc: Document) -> None:
    """Append an explicit page break paragraph."""
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(__import__("docx.enum.text", fromlist=["WD_BREAK"]).WD_BREAK.PAGE)


def _add_header_footer(doc: Document, company_name: str) -> None:
    """Add header with company name and footer with page numbers to all sections."""
    section = doc.sections[0]

    # --- Header ---
    section.different_first_page_header_footer = True
    header = section.header
    header.is_linked_to_previous = False
    para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = para.add_run(company_name)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)
    run.font.italic = True

    # Horizontal rule below header via bottom border on the paragraph
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)

    # --- Footer with page numbers ---
    footer = section.footer
    footer.is_linked_to_previous = False
    fpara = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fpara.clear()
    fpara.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Static prefix
    fpara.add_run("Page ").font.size = Pt(9)

    # PAGE field
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run_pg = fpara.add_run()
    run_pg.font.size = Pt(9)
    run_pg._r.append(fldChar1)
    run_pg._r.append(instrText)
    run_pg._r.append(fldChar2)

    fpara.add_run(" of ").font.size = Pt(9)

    # NUMPAGES field
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "begin")
    instrText2 = OxmlElement("w:instrText")
    instrText2.text = "NUMPAGES"
    fldChar4 = OxmlElement("w:fldChar")
    fldChar4.set(qn("w:fldCharType"), "end")
    run_np = fpara.add_run()
    run_np.font.size = Pt(9)
    run_np._r.append(fldChar3)
    run_np._r.append(instrText2)
    run_np._r.append(fldChar4)


# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------


def _ensure_style(doc: Document, style_name: str, base_style: str = "Normal") -> Any:
    """Return an existing style by name or create it from *base_style*."""
    try:
        return doc.styles[style_name]
    except KeyError:
        style = doc.styles.add_style(style_name, doc.styles[base_style].type)
        style.base_style = doc.styles[base_style]
        return style


def _configure_styles(doc: Document) -> None:
    """Apply consistent font and spacing to built-in and custom styles."""
    # Body text
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    # Heading levels
    for level, size, color_hex in [
        ("Heading 1", 20, "1F375D"),
        ("Heading 2", 14, "2C5282"),
        ("Heading 3", 12, "2B6CB0"),
    ]:
        try:
            style = doc.styles[level]
            style.font.name = "Calibri"
            style.font.size = Pt(size)
            style.font.color.rgb = RGBColor.from_string(color_hex)
            style.font.bold = True
        except KeyError:
            pass

    # Custom "Finding Title" style
    ft_style = _ensure_style(doc, "Finding Title", "Normal")
    ft_style.font.name = "Calibri"
    ft_style.font.size = Pt(12)
    ft_style.font.bold = True
    ft_style.font.color.rgb = RGBColor(31, 55, 93)


# ---------------------------------------------------------------------------
# Section builders
# ---------------------------------------------------------------------------


def _add_cover_page(doc: Document, data: ReportData) -> None:
    """Render a professional cover page."""
    # Large top spacer
    for _ in range(6):
        doc.add_paragraph()

    # Company / report type
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(data.company_name.upper())
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 100, 100)
    run.font.letter_spacing = Pt(2)

    # Report title
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Security Assessment Report")
    run2.font.name = "Calibri"
    run2.font.size = Pt(28)
    run2.font.bold = True
    run2.font.color.rgb = RGBColor(31, 55, 93)

    # Decorative separator
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("─" * 40)
    run3.font.color.rgb = RGBColor(200, 200, 200)

    doc.add_paragraph()

    # Client name
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run(data.client_name)
    run4.font.name = "Calibri"
    run4.font.size = Pt(18)
    run4.font.bold = True
    run4.font.color.rgb = RGBColor(44, 82, 130)

    # Target
    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run5 = p5.add_run(f"Target: {data.target}")
    run5.font.name = "Calibri"
    run5.font.size = Pt(12)
    run5.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph()

    # Metadata grid
    meta_lines = [
        ("Date:", data.scan_date),
        ("Scan ID:", data.scan_id),
        ("Classification:", "CONFIDENTIAL"),
    ]
    if data.author:
        meta_lines.insert(2, ("Prepared by:", data.author))

    tbl = doc.add_table(rows=len(meta_lines), cols=2)
    tbl.style = "Table Grid"
    for row_idx, (label, value) in enumerate(meta_lines):
        cells = tbl.rows[row_idx].cells
        cells[0].width = Inches(1.8)
        cells[1].width = Inches(3.5)
        cells[0].paragraphs[0].clear()
        run_l = cells[0].paragraphs[0].add_run(label)
        run_l.font.bold = True
        run_l.font.size = Pt(10)
        _set_cell_bg(cells[0], (240, 240, 240))
        cells[1].paragraphs[0].clear()
        run_v = cells[1].paragraphs[0].add_run(value)
        run_v.font.size = Pt(10)

    # Centre the table
    tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_executive_summary(doc: Document, data: ReportData) -> None:
    """Render executive summary section with severity breakdown table."""
    doc.add_heading("1. Executive Summary", level=1)

    # Severity summary table
    sev_counts = data.severity_counts
    tbl = doc.add_table(rows=1, cols=len(_SEVERITY_ORDER) + 1)
    tbl.style = "Table Grid"

    # Header row
    hdr_cells = tbl.rows[0].cells
    hdr_cells[0].paragraphs[0].clear()
    run = hdr_cells[0].paragraphs[0].add_run("Severity")
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    _set_cell_bg(hdr_cells[0], _HEADER_BG)

    for idx, sev in enumerate(_SEVERITY_ORDER):
        cell = hdr_cells[idx + 1]
        cell.paragraphs[0].clear()
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].add_run(sev.capitalize())
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        _set_cell_bg(cell, _HEADER_BG)

    # Count row
    count_row = tbl.add_row()
    count_cells = count_row.cells
    count_cells[0].paragraphs[0].clear()
    count_cells[0].paragraphs[0].add_run("Count").font.bold = True
    _set_cell_bg(count_cells[0], _ALT_ROW_BG)

    for idx, sev in enumerate(_SEVERITY_ORDER):
        cell = count_cells[idx + 1]
        cell.paragraphs[0].clear()
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        count = sev_counts.get(sev, 0)
        run = cell.paragraphs[0].add_run(str(count))
        rgb = _SEVERITY_COLORS.get(sev, (0, 0, 0))
        run.font.color.rgb = RGBColor(*rgb)
        run.font.bold = count > 0

    doc.add_paragraph()

    # Summary narrative
    summary_text = data.executive_summary or (
        f"This report presents the findings from a security assessment conducted against "
        f"{data.client_name} ({data.target}) on {data.scan_date}. "
        f"A total of {data.total_findings} finding(s) were identified across all severity levels. "
        f"The overall risk score is {data.risk_score:.1f} / 10.0."
    )
    para = doc.add_paragraph(summary_text)
    para.style = doc.styles["Normal"]


def _add_findings_table(doc: Document, data: ReportData) -> None:
    """Render the findings matrix table sorted by severity."""
    doc.add_heading("2. Table of Findings", level=1)

    if not data.findings:
        doc.add_paragraph("No findings were identified during this assessment.")
        return

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"

    # Header
    headers = ["#", "Title", "Severity", "Component"]
    hdr_row = tbl.rows[0]
    for idx, header_text in enumerate(headers):
        cell = hdr_row.cells[idx]
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(header_text)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        _set_cell_bg(cell, _HEADER_BG)
        _set_cell_borders(cell)

    # Rows grouped by severity
    row_num = 1
    for sev in _SEVERITY_ORDER:
        for finding in data.findings_by_severity.get(sev, []):
            data_row = tbl.add_row()
            cells = data_row.cells

            cells[0].paragraphs[0].clear()
            cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cells[0].paragraphs[0].add_run(str(row_num)).font.size = Pt(9)

            cells[1].paragraphs[0].clear()
            cells[1].paragraphs[0].add_run(finding.title).font.size = Pt(9)

            cells[2].paragraphs[0].clear()
            cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            sev_run = cells[2].paragraphs[0].add_run(
                finding.effective_severity.value.upper()
            )
            rgb = _SEVERITY_COLORS.get(finding.effective_severity.value, (0, 0, 0))
            sev_run.font.color.rgb = RGBColor(*rgb)
            sev_run.font.bold = True
            sev_run.font.size = Pt(9)

            cells[3].paragraphs[0].clear()
            cells[3].paragraphs[0].add_run(
                finding.affected_component or finding.target
            ).font.size = Pt(9)

            # Alternate row shading
            if row_num % 2 == 0:
                for cell in cells:
                    _set_cell_bg(cell, _ALT_ROW_BG)
            for cell in cells:
                _set_cell_borders(cell)

            row_num += 1


def _add_finding_detail(
    doc: Document,
    finding: Finding,
    index: int,
) -> None:
    """Render a single finding detail section."""
    # Finding heading
    h = doc.add_heading(level=2)
    h.clear()
    run_idx = h.add_run(f"{index}. ")
    run_idx.font.color.rgb = RGBColor(100, 100, 100)

    sev_rgb = _SEVERITY_COLORS.get(finding.effective_severity.value, (0, 0, 0))
    run_title = h.add_run(finding.title)
    run_title.font.color.rgb = RGBColor(31, 55, 93)

    run_badge = h.add_run(f"  [{finding.effective_severity.value.upper()}]")
    run_badge.font.color.rgb = RGBColor(*sev_rgb)
    run_badge.font.size = Pt(10)

    # Risk metadata table
    rows_data = [
        ("Severity", finding.effective_severity.value.upper()),
        ("Risk", finding.cvss.base_score if finding.cvss else "—"),
        ("CVSS Vector", finding.cvss.vector_string if finding.cvss else "—"),
        ("CVE IDs", ", ".join(finding.cve_ids) if finding.cve_ids else "—"),
        ("CWE IDs", ", ".join(finding.cwe_ids) if finding.cwe_ids else "—"),
        ("Target", finding.target),
        ("Component", finding.affected_component or "—"),
        ("Port", str(finding.port) if finding.port else "—"),
        ("Plugin", finding.source_plugin),
    ]

    tbl = doc.add_table(rows=len(rows_data), cols=2)
    tbl.style = "Table Grid"
    for row_idx, (label, value) in enumerate(rows_data):
        cells = tbl.rows[row_idx].cells
        cells[0].paragraphs[0].clear()
        lbl_run = cells[0].paragraphs[0].add_run(label)
        lbl_run.font.bold = True
        lbl_run.font.size = Pt(9)
        _set_cell_bg(cells[0], _ALT_ROW_BG)
        _set_cell_borders(cells[0])

        cells[1].paragraphs[0].clear()
        val_run = cells[1].paragraphs[0].add_run(str(value))
        val_run.font.size = Pt(9)
        _set_cell_borders(cells[1])

    doc.add_paragraph()

    # Description
    doc.add_heading("Description", level=3)
    desc_para = doc.add_paragraph(finding.description)
    desc_para.style = doc.styles["Normal"]

    # Evidence
    if finding.evidence:
        doc.add_heading("Evidence", level=3)
        for ev in finding.evidence:
            p = doc.add_paragraph(style="Normal")
            run_ev_title = p.add_run(f"{ev.title}: ")
            run_ev_title.font.bold = True
            p.add_run(ev.content)

    # Recommendation
    if finding.remediation:
        doc.add_heading("Recommendation", level=3)
        doc.add_paragraph(finding.remediation, style="Normal")

    # References
    if finding.references:
        doc.add_heading("References", level=3)
        for ref in finding.references:
            ref_para = doc.add_paragraph(style="List Bullet")
            ref_run = ref_para.add_run(ref.title)
            ref_run.font.bold = True
            ref_para.add_run(f" — {ref.url}")


def _add_appendix(doc: Document, data: ReportData) -> None:
    """Render the methodology and tool versions appendix."""
    doc.add_heading("Appendix A — Methodology", level=1)
    doc.add_paragraph(data.methodology, style="Normal")

    doc.add_heading("Appendix B — Tool Versions", level=2)
    p = doc.add_paragraph(
        "Automated scanning was performed using the following tools:",
        style="Normal",
    )
    tools = [
        "VXIS Platform (current version)",
        "Nuclei — Template-based vulnerability scanner",
        "Nmap — Network discovery and security auditing",
        "TestSSL.sh — TLS/SSL configuration tester",
        "TruffleHog — Secret detection in source code",
        "CheckDMARC — Email authentication record validation",
        "WAFW00F — Web Application Firewall detection",
    ]
    for tool in tools:
        doc.add_paragraph(tool, style="List Bullet")


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


class DOCXReportGenerator:
    """Generate professional DOCX reports from VXIS scan results.

    Produces NCC Group-style Word documents with:
    - Cover page
    - Executive summary with severity breakdown table
    - Table of findings matrix
    - Per-finding detail sections
    - Methodology appendix
    """

    def generate(self, data: ReportData, output_path: Path) -> Path:
        """Create and write a DOCX report to *output_path*.

        Parameters
        ----------
        data:
            Populated :class:`ReportData` instance from the report generator.
        output_path:
            Destination file path for the generated .docx file.

        Returns
        -------
        Path
            Resolved path of the written file.
        """
        output_path = Path(output_path).resolve()
        output_path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()
        _configure_styles(doc)
        _add_header_footer(doc, data.company_name)

        # 1. Cover page
        _add_cover_page(doc, data)
        _add_page_break(doc)

        # 2. Executive summary
        _add_executive_summary(doc, data)
        _add_page_break(doc)

        # 3. Table of findings
        _add_findings_table(doc, data)
        _add_page_break(doc)

        # 4. Finding details
        doc.add_heading("3. Finding Details", level=1)

        if not data.findings:
            doc.add_paragraph(
                "No findings were identified during this assessment.",
                style="Normal",
            )
        else:
            finding_index = 1
            for sev in _SEVERITY_ORDER:
                for finding in data.findings_by_severity.get(sev, []):
                    _add_finding_detail(doc, finding, finding_index)
                    finding_index += 1
                    doc.add_paragraph()

        _add_page_break(doc)

        # 5. Appendix
        _add_appendix(doc, data)

        doc.save(str(output_path))
        return output_path
