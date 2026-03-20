"""Attestation letter generator for VXIS security assessments.

Produces a single-page DOCX letter suitable for compliance or due-diligence
purposes, summarising the scope and severity distribution of a completed
security assessment.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from vxis.report.generator import ReportData

_SEVERITY_ORDER: list[str] = [
    "critical",
    "high",
    "medium",
    "low",
    "informational",
]

_SEVERITY_COLORS: dict[str, tuple[int, int, int]] = {
    "critical": (123, 44, 52),
    "high": (192, 57, 43),
    "medium": (230, 126, 34),
    "low": (46, 204, 113),
    "informational": (52, 152, 219),
}

_HEADER_BG: tuple[int, int, int] = (31, 55, 93)
_ALT_ROW_BG: tuple[int, int, int] = (242, 242, 242)

_DISCLAIMER = (
    "This attestation letter is issued solely for informational purposes and is "
    "based on the results of the automated and manual security assessment conducted "
    "by {company} during the period referenced above. The assessment was performed "
    "using industry-standard tools and methodologies; however, no security assessment "
    "can guarantee the complete absence of vulnerabilities. This letter should not be "
    "construed as a warranty or guarantee of security. {company} accepts no liability "
    "for damages arising from reliance on this document."
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _set_cell_bg(cell, rgb: tuple[int, int, int]) -> None:
    hex_color = "{:02X}{:02X}{:02X}".format(*rgb)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_borders(cell) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "CCCCCC")
        tcBorders.append(border)
    tcPr.append(tcBorders)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


class AttestationGenerator:
    """Generate a one-page attestation letter in DOCX format.

    The letter includes:
    - Company letterhead
    - Assessment date and classification
    - Scope paragraph
    - Severity count table per component
    - Disclaimer paragraph
    - Signature block
    """

    def generate(self, data: ReportData, output_path: Path) -> Path:
        """Create and write the attestation letter to *output_path*.

        Parameters
        ----------
        data:
            Populated :class:`ReportData` instance.
        output_path:
            Destination file path for the generated .docx file.

        Returns
        -------
        Path
            Resolved path of the written file.
        """
        output_path = Path(output_path).resolve()
        output_path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()
        self._configure_page(doc)
        self._add_letterhead(doc, data)
        self._add_salutation(doc, data)
        self._add_scope_paragraph(doc, data)
        self._add_severity_table(doc, data)
        self._add_disclaimer(doc, data)
        self._add_signature_block(doc, data)

        doc.save(str(output_path))
        return output_path

    # ------------------------------------------------------------------
    # Section builders
    # ------------------------------------------------------------------

    def _configure_page(self, doc: Document) -> None:
        """Set narrow margins so content fits on one page."""
        section = doc.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)

    def _add_letterhead(self, doc: Document, data: ReportData) -> None:
        """Render company name and address block at the top."""
        # Company name — large and bold
        p_company = doc.add_paragraph()
        p_company.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p_company.add_run(data.company_name)
        run.font.name = "Calibri"
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(31, 55, 93)

        # Decorative rule
        p_rule = doc.add_paragraph()
        p_rule.alignment = WD_ALIGN_PARAGRAPH.LEFT
        rule_run = p_rule.add_run("─" * 60)
        rule_run.font.color.rgb = RGBColor(200, 200, 200)
        rule_run.font.size = Pt(8)

        # Date and classification on same line (right-aligned date)
        p_meta = doc.add_paragraph()
        p_meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        meta_run = p_meta.add_run(
            f"Date: {data.scan_date}    |    Classification: CONFIDENTIAL"
        )
        meta_run.font.size = Pt(9)
        meta_run.font.color.rgb = RGBColor(100, 100, 100)
        meta_run.font.italic = True

        doc.add_paragraph()

    def _add_salutation(self, doc: Document, data: ReportData) -> None:
        """Add salutation line."""
        p = doc.add_paragraph()
        p.add_run("To Whom It May Concern,").font.bold = True
        doc.add_paragraph()

    def _add_scope_paragraph(self, doc: Document, data: ReportData) -> None:
        """Render the scope description paragraph."""
        counts = data.severity_counts
        total = data.total_findings
        critical = counts.get("critical", 0)
        high = counts.get("high", 0)
        medium = counts.get("medium", 0)
        low = counts.get("low", 0)
        info = counts.get("informational", 0)

        scope_text = (
            f"{data.company_name} has conducted a security assessment of the systems "
            f"and infrastructure belonging to {data.client_name} with primary target "
            f"'{data.target}'. The assessment was carried out on {data.scan_date} "
            f"(Scan reference: {data.scan_id}) using automated and manual testing "
            f"methodologies in accordance with industry standards including OWASP and PTES.\n\n"
            f"A total of {total} security finding(s) were identified and classified as follows: "
            f"{critical} Critical, {high} High, {medium} Medium, {low} Low, "
            f"and {info} Informational."
        )
        p = doc.add_paragraph(scope_text)
        p.style = doc.styles["Normal"]
        doc.add_paragraph()

    def _add_severity_table(self, doc: Document, data: ReportData) -> None:
        """Render severity count table.

        Columns: Severity | Count
        One row per severity level.
        """
        doc.add_paragraph("Severity Summary:").runs[0].font.bold = True

        counts = data.severity_counts
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"

        # Header
        hdr_cells = tbl.rows[0].cells
        for idx, label in enumerate(["Severity Level", "Finding Count"]):
            hdr_cells[idx].paragraphs[0].clear()
            run = hdr_cells[idx].paragraphs[0].add_run(label)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(9)
            _set_cell_bg(hdr_cells[idx], _HEADER_BG)
            _set_cell_borders(hdr_cells[idx])

        # Data rows
        for row_idx, sev in enumerate(_SEVERITY_ORDER):
            row = tbl.add_row()
            cells = row.cells

            cells[0].paragraphs[0].clear()
            sev_run = cells[0].paragraphs[0].add_run(sev.capitalize())
            rgb = _SEVERITY_COLORS.get(sev, (0, 0, 0))
            sev_run.font.color.rgb = RGBColor(*rgb)
            sev_run.font.bold = True
            sev_run.font.size = Pt(9)

            cells[1].paragraphs[0].clear()
            cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            count_run = cells[1].paragraphs[0].add_run(str(counts.get(sev, 0)))
            count_run.font.size = Pt(9)

            if row_idx % 2 == 0:
                _set_cell_bg(cells[0], _ALT_ROW_BG)
                _set_cell_bg(cells[1], _ALT_ROW_BG)

            _set_cell_borders(cells[0])
            _set_cell_borders(cells[1])

        doc.add_paragraph()

    def _add_disclaimer(self, doc: Document, data: ReportData) -> None:
        """Render the legal disclaimer paragraph."""
        doc.add_paragraph("Disclaimer:").runs[0].font.bold = True
        disclaimer = _DISCLAIMER.format(company=data.company_name)
        p = doc.add_paragraph(disclaimer)
        p.style = doc.styles["Normal"]
        for run in p.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(80, 80, 80)
        doc.add_paragraph()

    def _add_signature_block(self, doc: Document, data: ReportData) -> None:
        """Render signature block at the bottom of the letter."""
        p_sig = doc.add_paragraph()
        p_sig.add_run("Sincerely,\n").font.bold = False
        run_co = p_sig.add_run(data.company_name)
        run_co.font.bold = True
        run_co.font.color.rgb = RGBColor(31, 55, 93)

        if data.author:
            doc.add_paragraph()
            p_auth = doc.add_paragraph()
            p_auth.add_run(data.author).font.bold = True

        # Signature line
        doc.add_paragraph()
        p_line = doc.add_paragraph()
        line_run = p_line.add_run("_" * 30)
        line_run.font.color.rgb = RGBColor(150, 150, 150)

        p_role = doc.add_paragraph()
        role_run = p_role.add_run(
            f"Authorized Representative, {data.company_name}"
        )
        role_run.font.size = Pt(9)
        role_run.font.color.rgb = RGBColor(100, 100, 100)
        role_run.font.italic = True
