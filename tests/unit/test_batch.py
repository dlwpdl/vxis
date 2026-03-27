"""Unit tests for VXIS PE portfolio batch scanning.

Covers BatchScanner, BatchTarget, BatchResult, and the summary report
generator without actually hitting the network or requiring a running
scan environment.
"""

from __future__ import annotations

import io
import textwrap
from datetime import datetime, timezone
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest

pytest.importorskip("docx", reason="python-docx not installed (optional 'export' extra)")

from docx import Document  # noqa: E402

from vxis.core.batch import BatchResult, BatchScanner, BatchTarget  # noqa: E402
from vxis.core.orchestrator import ScanResult  # noqa: E402
from vxis.models.finding import Finding, Severity  # noqa: E402


# ---------------------------------------------------------------------------
# Shared helpers
# ---------------------------------------------------------------------------


def make_scan_result(
    scan_id: str = "scan-001",
    target: str = "acme.com",
    profile: str = "standard",
    findings: list[Finding] | None = None,
) -> ScanResult:
    """Return a minimal ScanResult with sensible defaults."""
    return ScanResult(
        scan_id=scan_id,
        target=target,
        profile=profile,
        findings=findings or [],
        tool_runs=[],
        started_at=datetime.now(timezone.utc),
        finished_at=datetime.now(timezone.utc),
    )


def make_finding(
    id: str = "f-001",
    title: str = "SQL Injection",
    severity: Severity = Severity.high,
    finding_type: str = "sqli",
    **overrides,
) -> Finding:
    """Return a minimal valid Finding."""
    defaults: dict = dict(
        id=id,
        scan_id="scan-001",
        title=title,
        description="Test finding description.",
        severity=severity,
        target="acme.com",
        finding_type=finding_type,
        source_plugin="test_plugin",
    )
    defaults.update(overrides)
    return Finding(**defaults)


def make_batch_target(
    name: str = "ACME Corp",
    domain: str = "acme.com",
    extra_domains: list[str] | None = None,
    notes: str = "PE portfolio company",
) -> BatchTarget:
    return BatchTarget(
        name=name,
        domain=domain,
        extra_domains=extra_domains or [],
        notes=notes,
    )


CSV_CONTENT = textwrap.dedent("""\
    name,domain,extra_domains,notes
    "ACME Corp",acme.com,"api.acme.com;staging.acme.com","PE portfolio company"
    "Beta Inc",beta.io,,""
    "Gamma LLC",gamma.co,"app.gamma.co","SaaS tool"
""")


# ---------------------------------------------------------------------------
# BatchTarget — load_targets
# ---------------------------------------------------------------------------


class TestLoadTargets:
    def test_load_targets_from_csv_string(self) -> None:
        """load_targets must parse a StringIO as if it were a file."""
        csv_io = io.StringIO(CSV_CONTENT)
        targets = BatchScanner.load_targets(csv_io)  # type: ignore[arg-type]

        assert len(targets) == 3

    def test_load_targets_names(self) -> None:
        """Company names must be correctly parsed."""
        csv_io = io.StringIO(CSV_CONTENT)
        targets = BatchScanner.load_targets(csv_io)  # type: ignore[arg-type]

        assert targets[0].name == "ACME Corp"
        assert targets[1].name == "Beta Inc"
        assert targets[2].name == "Gamma LLC"

    def test_load_targets_domains(self) -> None:
        """Primary domains must be correctly parsed."""
        csv_io = io.StringIO(CSV_CONTENT)
        targets = BatchScanner.load_targets(csv_io)  # type: ignore[arg-type]

        assert targets[0].domain == "acme.com"
        assert targets[1].domain == "beta.io"
        assert targets[2].domain == "gamma.co"

    def test_load_targets_extra_domains_semicolon_separation(self) -> None:
        """extra_domains column must be split on semicolons."""
        csv_io = io.StringIO(CSV_CONTENT)
        targets = BatchScanner.load_targets(csv_io)  # type: ignore[arg-type]

        assert targets[0].extra_domains == ["api.acme.com", "staging.acme.com"]

    def test_load_targets_empty_extra_domains(self) -> None:
        """Empty extra_domains column must produce an empty list."""
        csv_io = io.StringIO(CSV_CONTENT)
        targets = BatchScanner.load_targets(csv_io)  # type: ignore[arg-type]

        assert targets[1].extra_domains == []

    def test_load_targets_notes(self) -> None:
        """Notes column must be correctly parsed."""
        csv_io = io.StringIO(CSV_CONTENT)
        targets = BatchScanner.load_targets(csv_io)  # type: ignore[arg-type]

        assert targets[0].notes == "PE portfolio company"
        assert targets[2].notes == "SaaS tool"

    def test_load_targets_single_extra_domain(self) -> None:
        """A single extra domain (no semicolon) must produce a list of one."""
        csv_io = io.StringIO(CSV_CONTENT)
        targets = BatchScanner.load_targets(csv_io)  # type: ignore[arg-type]

        assert targets[2].extra_domains == ["app.gamma.co"]

    def test_load_targets_from_real_file(self, tmp_path: Path) -> None:
        """load_targets must work with a real file path."""
        csv_file = tmp_path / "portfolio.csv"
        csv_file.write_text(CSV_CONTENT, encoding="utf-8")

        targets = BatchScanner.load_targets(csv_file)

        assert len(targets) == 3
        assert targets[0].domain == "acme.com"


# ---------------------------------------------------------------------------
# BatchScanner — calculate_risk_grade
# ---------------------------------------------------------------------------


class TestCalculateRiskGrade:
    def test_grade_a_no_critical_no_high(self) -> None:
        """0 critical, 0 high → grade A."""
        sr = make_scan_result(findings=[
            make_finding(id="f1", severity=Severity.medium),
            make_finding(id="f2", severity=Severity.low),
        ])
        assert BatchScanner.calculate_risk_grade(sr) == "A"

    def test_grade_a_empty_findings(self) -> None:
        """No findings at all → grade A."""
        sr = make_scan_result(findings=[])
        assert BatchScanner.calculate_risk_grade(sr) == "A"

    def test_grade_b_one_high(self) -> None:
        """0 critical, 1 high → grade B."""
        sr = make_scan_result(findings=[
            make_finding(id="f1", severity=Severity.high),
        ])
        assert BatchScanner.calculate_risk_grade(sr) == "B"

    def test_grade_b_three_high(self) -> None:
        """0 critical, 3 high → grade B."""
        sr = make_scan_result(findings=[
            make_finding(id="f1", severity=Severity.high),
            make_finding(id="f2", severity=Severity.high),
            make_finding(id="f3", severity=Severity.high),
        ])
        assert BatchScanner.calculate_risk_grade(sr) == "B"

    def test_grade_c_four_high(self) -> None:
        """0 critical, 4 high → grade C."""
        sr = make_scan_result(findings=[
            make_finding(id=f"f{i}", severity=Severity.high) for i in range(4)
        ])
        assert BatchScanner.calculate_risk_grade(sr) == "C"

    def test_grade_c_one_critical(self) -> None:
        """1 critical → grade C."""
        sr = make_scan_result(findings=[
            make_finding(id="f1", severity=Severity.critical),
        ])
        assert BatchScanner.calculate_risk_grade(sr) == "C"

    def test_grade_d_two_critical(self) -> None:
        """2 critical → grade D."""
        sr = make_scan_result(findings=[
            make_finding(id="f1", severity=Severity.critical),
            make_finding(id="f2", severity=Severity.critical),
        ])
        assert BatchScanner.calculate_risk_grade(sr) == "D"

    def test_grade_d_three_critical(self) -> None:
        """3 critical → grade D."""
        sr = make_scan_result(findings=[
            make_finding(id=f"f{i}", severity=Severity.critical) for i in range(3)
        ])
        assert BatchScanner.calculate_risk_grade(sr) == "D"

    def test_grade_f_four_critical(self) -> None:
        """4 critical → grade F."""
        sr = make_scan_result(findings=[
            make_finding(id=f"f{i}", severity=Severity.critical) for i in range(4)
        ])
        assert BatchScanner.calculate_risk_grade(sr) == "F"

    def test_grade_f_many_critical(self) -> None:
        """10 critical → grade F."""
        sr = make_scan_result(findings=[
            make_finding(id=f"f{i}", severity=Severity.critical) for i in range(10)
        ])
        assert BatchScanner.calculate_risk_grade(sr) == "F"


# ---------------------------------------------------------------------------
# BatchResult
# ---------------------------------------------------------------------------


class TestBatchResult:
    def test_batch_result_succeeded_when_no_error(self) -> None:
        """BatchResult.succeeded is True when scan_result is set and error is None."""
        target = make_batch_target()
        sr = make_scan_result()
        result = BatchResult(target=target, scan_result=sr)

        assert result.succeeded is True

    def test_batch_result_failed_when_error_set(self) -> None:
        """BatchResult.succeeded is False when error is set."""
        target = make_batch_target()
        result = BatchResult(target=target, scan_result=None, error="connection refused")

        assert result.succeeded is False

    def test_batch_result_failed_when_scan_result_none(self) -> None:
        """BatchResult.succeeded is False when scan_result is None (no error message)."""
        target = make_batch_target()
        result = BatchResult(target=target, scan_result=None)

        assert result.succeeded is False

    def test_batch_result_preserves_target(self) -> None:
        """BatchResult must preserve the target reference."""
        target = make_batch_target(name="SpecialCo", domain="special.co")
        result = BatchResult(target=target, scan_result=None, error="timeout")

        assert result.target.name == "SpecialCo"
        assert result.target.domain == "special.co"
        assert result.error == "timeout"

    def test_batch_result_with_error_string(self) -> None:
        """Error field must be a plain string."""
        target = make_batch_target()
        result = BatchResult(
            target=target,
            scan_result=None,
            error="network unreachable: ETIMEDOUT",
        )

        assert result.error == "network unreachable: ETIMEDOUT"


# ---------------------------------------------------------------------------
# BatchScanner — generate_summary_report
# ---------------------------------------------------------------------------


class TestGenerateSummaryReport:
    def _make_scanner(self) -> BatchScanner:
        """Return a BatchScanner with a mocked config."""
        config = MagicMock()
        config.profiles = {"standard": MagicMock(skip_plugins=[], tool_overrides={}, max_concurrency=4)}
        config.data_dir = Path("/tmp/vxis_test")
        config.db_url = "sqlite+aiosqlite:////tmp/vxis_test.db"
        with patch("vxis.core.batch.ScanOrchestrator"):
            scanner = BatchScanner(config)
        return scanner

    def test_generate_summary_report_creates_file(self, tmp_path: Path) -> None:
        """generate_summary_report must create a .docx file at the given path."""
        scanner = self._make_scanner()

        target1 = make_batch_target(name="ACME Corp", domain="acme.com")
        target2 = make_batch_target(name="Beta Inc", domain="beta.io")

        sr1 = make_scan_result(
            target="acme.com",
            findings=[
                make_finding(id="f1", severity=Severity.critical),
                make_finding(id="f2", severity=Severity.high),
            ],
        )
        sr2 = make_scan_result(
            target="beta.io",
            findings=[],
        )

        results = [
            BatchResult(target=target1, scan_result=sr1),
            BatchResult(target=target2, scan_result=sr2),
        ]

        out = tmp_path / "summary.docx"
        returned = scanner.generate_summary_report(results, out)

        assert returned == out
        assert out.exists()
        assert out.stat().st_size > 0

    def test_generate_summary_report_contains_company_names(self, tmp_path: Path) -> None:
        """Summary report must contain the names of all scanned companies."""
        scanner = self._make_scanner()

        results = [
            BatchResult(
                target=make_batch_target(name="AlphaCo", domain="alpha.com"),
                scan_result=make_scan_result(target="alpha.com"),
            ),
            BatchResult(
                target=make_batch_target(name="BetaCorp", domain="beta.com"),
                scan_result=None,
                error="scan failed",
            ),
        ]

        out = tmp_path / "summary.docx"
        scanner.generate_summary_report(results, out)

        doc = Document(str(out))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        all_table_text = "\n".join(
            cell.text for table in doc.tables for row in table.rows for cell in row.cells
        )
        full_text = all_text + all_table_text

        assert "AlphaCo" in full_text
        assert "BetaCorp" in full_text

    def test_generate_summary_report_grade_appears(self, tmp_path: Path) -> None:
        """Summary report risk table must contain grade letters."""
        scanner = self._make_scanner()

        sr = make_scan_result(
            findings=[make_finding(id="f1", severity=Severity.medium)],
        )
        results = [
            BatchResult(
                target=make_batch_target(),
                scan_result=sr,
            )
        ]

        out = tmp_path / "summary.docx"
        scanner.generate_summary_report(results, out)

        doc = Document(str(out))
        all_table_text = "\n".join(
            cell.text for table in doc.tables for row in table.rows for cell in row.cells
        )

        # Grade A expected since no high or critical
        assert "A" in all_table_text

    def test_generate_summary_report_empty_results(self, tmp_path: Path) -> None:
        """generate_summary_report must not crash on an empty results list."""
        scanner = self._make_scanner()

        out = tmp_path / "empty_summary.docx"
        scanner.generate_summary_report([], out)

        assert out.exists()

    def test_generate_summary_creates_parent_dirs(self, tmp_path: Path) -> None:
        """generate_summary_report must create intermediate directories."""
        scanner = self._make_scanner()
        out = tmp_path / "a" / "b" / "c" / "summary.docx"

        scanner.generate_summary_report([], out)

        assert out.exists()
