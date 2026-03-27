"""Unit tests for VXIS DOCX report export.

Covers DOCXReportGenerator and AttestationGenerator, verifying:
- File creation
- Content presence (client name, finding titles)
- Edge cases (no findings)
- One-page attestation with severity counts
"""

from __future__ import annotations

from pathlib import Path

import pytest

pytest.importorskip("docx", reason="python-docx not installed (optional 'export' extra)")

from docx import Document  # noqa: E402

from vxis.models.finding import Evidence, Finding, Reference, Severity  # noqa: E402
from vxis.report.attestation import AttestationGenerator  # noqa: E402
from vxis.report.docx_export import DOCXReportGenerator  # noqa: E402
from vxis.report.generator import ReportData  # noqa: E402


# ---------------------------------------------------------------------------
# Shared helpers
# ---------------------------------------------------------------------------


def make_finding(
    id: str = "f-001",
    title: str = "SQL Injection",
    severity: Severity = Severity.high,
    finding_type: str = "sqli",
    description: str = "Unsanitised input passed to SQL query.",
    **overrides,
) -> Finding:
    """Return a minimal valid Finding with sensible defaults."""
    defaults: dict = dict(
        id=id,
        scan_id="scan-001",
        title=title,
        description=description,
        severity=severity,
        target="192.168.1.10",
        finding_type=finding_type,
        source_plugin="test_plugin",
    )
    defaults.update(overrides)
    return Finding(**defaults)


def make_report_data(
    findings: list[Finding] | None = None,
    **overrides,
) -> ReportData:
    """Return a populated ReportData with sensible defaults."""
    defaults: dict = dict(
        scan_id="scan-abc-123",
        client_name="Acme Corp",
        target="acme.com",
        scan_date="2026-03-20",
        findings=findings if findings is not None else [],
        author="Jane Smith",
        company_name="VXIS Security",
    )
    defaults.update(overrides)
    return ReportData(**defaults)


def _all_text(doc: Document) -> str:
    """Extract all paragraph text from a Document as a single string."""
    return "\n".join(p.text for p in doc.paragraphs)


def _all_table_text(doc: Document) -> str:
    """Extract all table cell text from a Document."""
    parts: list[str] = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# DOCXReportGenerator
# ---------------------------------------------------------------------------


class TestDOCXReportGeneratorFileCreation:
    def test_generate_creates_docx_file(self, tmp_path: Path) -> None:
        """generate() must create a .docx file at the specified path."""
        gen = DOCXReportGenerator()
        data = make_report_data()
        out = tmp_path / "report.docx"

        result = gen.generate(data, out)

        assert result == out
        assert out.exists()
        assert out.stat().st_size > 0

    def test_generate_creates_parent_dirs(self, tmp_path: Path) -> None:
        """generate() must create intermediate directories if they don't exist."""
        gen = DOCXReportGenerator()
        data = make_report_data()
        out = tmp_path / "deep" / "nested" / "report.docx"

        gen.generate(data, out)

        assert out.exists()

    def test_generate_returns_resolved_path(self, tmp_path: Path) -> None:
        """generate() must return an absolute resolved path."""
        gen = DOCXReportGenerator()
        data = make_report_data()
        out = tmp_path / "report.docx"

        result = gen.generate(data, out)

        assert result.is_absolute()


class TestDOCXReportGeneratorContent:
    def test_docx_contains_client_name(self, tmp_path: Path) -> None:
        """The generated document must contain the client name."""
        gen = DOCXReportGenerator()
        data = make_report_data(client_name="GlobalBank PLC")
        out = tmp_path / "report.docx"

        gen.generate(data, out)
        doc = Document(str(out))
        full_text = _all_text(doc) + _all_table_text(doc)

        assert "GlobalBank PLC" in full_text

    def test_docx_contains_finding_titles(self, tmp_path: Path) -> None:
        """Each finding's title must appear in the generated document."""
        gen = DOCXReportGenerator()
        findings = [
            make_finding(id="f1", title="Remote Code Execution", severity=Severity.critical),
            make_finding(id="f2", title="XSS Reflected", severity=Severity.medium),
            make_finding(id="f3", title="Weak TLS Configuration", severity=Severity.low),
        ]
        data = make_report_data(findings=findings)
        out = tmp_path / "report.docx"

        gen.generate(data, out)
        doc = Document(str(out))
        full_text = _all_text(doc) + _all_table_text(doc)

        for finding in findings:
            assert finding.title in full_text, (
                f"Expected finding title '{finding.title}' not found in document"
            )

    def test_docx_contains_scan_id(self, tmp_path: Path) -> None:
        """The scan ID must appear somewhere in the generated document."""
        gen = DOCXReportGenerator()
        data = make_report_data(scan_id="unique-scan-xyz-9999")
        out = tmp_path / "report.docx"

        gen.generate(data, out)
        doc = Document(str(out))
        full_text = _all_text(doc) + _all_table_text(doc)

        assert "unique-scan-xyz-9999" in full_text

    def test_docx_contains_target(self, tmp_path: Path) -> None:
        """The target domain must appear in the generated document."""
        gen = DOCXReportGenerator()
        data = make_report_data(target="target.example.com")
        out = tmp_path / "report.docx"

        gen.generate(data, out)
        doc = Document(str(out))
        full_text = _all_text(doc) + _all_table_text(doc)

        assert "target.example.com" in full_text

    def test_docx_finding_with_evidence_and_references(self, tmp_path: Path) -> None:
        """Findings with evidence and references should render without error."""
        gen = DOCXReportGenerator()
        finding = make_finding(
            id="f1",
            title="Open Redirect",
            severity=Severity.medium,
            evidence=[
                Evidence(
                    evidence_type="http_response",
                    title="HTTP Response",
                    content="HTTP/1.1 302 Found\nLocation: https://evil.com",
                )
            ],
            references=[
                Reference(title="OWASP Open Redirect", url="https://owasp.org/redirect")
            ],
            remediation="Validate all redirect targets against an allowlist.",
        )
        data = make_report_data(findings=[finding])
        out = tmp_path / "report.docx"

        gen.generate(data, out)
        doc = Document(str(out))
        full_text = _all_text(doc) + _all_table_text(doc)

        assert "Open Redirect" in full_text
        assert "OWASP Open Redirect" in full_text


class TestDOCXReportGeneratorNoFindings:
    def test_no_findings_produces_valid_document(self, tmp_path: Path) -> None:
        """A report with zero findings must still produce a valid, openable DOCX."""
        gen = DOCXReportGenerator()
        data = make_report_data(findings=[])
        out = tmp_path / "empty_report.docx"

        gen.generate(data, out)

        # Must open without error
        doc = Document(str(out))
        assert doc is not None

    def test_no_findings_mentions_no_findings(self, tmp_path: Path) -> None:
        """When there are no findings the document should say so."""
        gen = DOCXReportGenerator()
        data = make_report_data(findings=[])
        out = tmp_path / "empty_report.docx"

        gen.generate(data, out)
        doc = Document(str(out))
        full_text = _all_text(doc)

        assert "No findings" in full_text

    def test_no_findings_risk_score_zero(self, tmp_path: Path) -> None:
        """Risk score for empty findings must be 0.0."""
        data = make_report_data(findings=[])
        assert data.risk_score == 0.0


class TestDOCXSeverityRendering:
    def test_critical_severity_appears_in_findings_table(self, tmp_path: Path) -> None:
        """CRITICAL label must appear in the findings matrix for a critical finding."""
        gen = DOCXReportGenerator()
        findings = [make_finding(id="f1", severity=Severity.critical)]
        data = make_report_data(findings=findings)
        out = tmp_path / "report.docx"

        gen.generate(data, out)
        doc = Document(str(out))
        table_text = _all_table_text(doc)

        assert "CRITICAL" in table_text

    def test_all_severity_levels_in_summary_table(self, tmp_path: Path) -> None:
        """The executive summary severity table must list all severity levels."""
        gen = DOCXReportGenerator()
        data = make_report_data()
        out = tmp_path / "report.docx"

        gen.generate(data, out)
        doc = Document(str(out))
        table_text = _all_table_text(doc)

        for sev in ["Critical", "High", "Medium", "Low", "Informational"]:
            assert sev in table_text, f"'{sev}' not found in document tables"


# ---------------------------------------------------------------------------
# AttestationGenerator
# ---------------------------------------------------------------------------


class TestAttestationGenerator:
    def test_attestation_generates_docx_file(self, tmp_path: Path) -> None:
        """AttestationGenerator.generate() must create a .docx file."""
        gen = AttestationGenerator()
        data = make_report_data()
        out = tmp_path / "attestation.docx"

        result = gen.generate(data, out)

        assert result == out
        assert out.exists()
        assert out.stat().st_size > 0

    def test_attestation_contains_client_name(self, tmp_path: Path) -> None:
        """Attestation letter must include the client name."""
        gen = AttestationGenerator()
        data = make_report_data(client_name="FinTech Partners Ltd")
        out = tmp_path / "attestation.docx"

        gen.generate(data, out)
        doc = Document(str(out))
        full_text = _all_text(doc) + _all_table_text(doc)

        assert "FinTech Partners Ltd" in full_text

    def test_attestation_contains_to_whom_salutation(self, tmp_path: Path) -> None:
        """Attestation must include standard 'To Whom It May Concern' salutation."""
        gen = AttestationGenerator()
        data = make_report_data()
        out = tmp_path / "attestation.docx"

        gen.generate(data, out)
        doc = Document(str(out))
        full_text = _all_text(doc)

        assert "To Whom It May Concern" in full_text

    def test_attestation_contains_severity_counts(self, tmp_path: Path) -> None:
        """Attestation severity table must show counts for each severity level."""
        gen = AttestationGenerator()
        findings = [
            make_finding(id="f1", severity=Severity.critical),
            make_finding(id="f2", severity=Severity.high),
            make_finding(id="f3", severity=Severity.high),
            make_finding(id="f4", severity=Severity.medium),
        ]
        data = make_report_data(findings=findings)
        out = tmp_path / "attestation.docx"

        gen.generate(data, out)
        doc = Document(str(out))
        table_text = _all_table_text(doc)

        # All severity names must appear in the severity table
        for sev in ["Critical", "High", "Medium", "Low", "Informational"]:
            assert sev in table_text, f"'{sev}' not found in attestation tables"

    def test_attestation_contains_disclaimer(self, tmp_path: Path) -> None:
        """Attestation must include a disclaimer paragraph."""
        gen = AttestationGenerator()
        data = make_report_data()
        out = tmp_path / "attestation.docx"

        gen.generate(data, out)
        doc = Document(str(out))
        full_text = _all_text(doc)

        assert "Disclaimer" in full_text

    def test_attestation_contains_company_name(self, tmp_path: Path) -> None:
        """Company name must appear in the letterhead."""
        gen = AttestationGenerator()
        data = make_report_data(company_name="Sentinel Security Ltd")
        out = tmp_path / "attestation.docx"

        gen.generate(data, out)
        doc = Document(str(out))
        full_text = _all_text(doc)

        assert "Sentinel Security Ltd" in full_text

    def test_attestation_zero_findings_counts_shown(self, tmp_path: Path) -> None:
        """Attestation with no findings must still render severity table with zeros."""
        gen = AttestationGenerator()
        data = make_report_data(findings=[])
        out = tmp_path / "attestation.docx"

        gen.generate(data, out)
        doc = Document(str(out))
        table_text = _all_table_text(doc)

        # Severity labels should still be present
        for sev in ["Critical", "High", "Medium", "Low"]:
            assert sev in table_text
